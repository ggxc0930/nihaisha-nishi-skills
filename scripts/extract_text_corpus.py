#!/usr/bin/env python3
"""Extract local text corpus for later Skill distillation."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
import re


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PLAN = ROOT / "distillation-output" / "distillation-plan.jsonl"
DEFAULT_OUTPUT = ROOT / "distillation-output" / "extracted-text"


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path, max_pages: int | None) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        pages = pdf.pages[:max_pages] if max_pages else pdf.pages
        for index, page in enumerate(pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n\n## Page {index}\n\n{text}")
    return clean_text("\n".join(parts))


def extract_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def extract_xlsx(path: Path, max_rows: int) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"\n\n## Sheet: {sheet.title}\n")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_index > max_rows:
                parts.append("[truncated]")
                break
            values = [str(value) for value in row if value is not None and str(value).strip()]
            if values:
                parts.append(" | ".join(values))
    return clean_text("\n".join(parts))


def extract_plain(path: Path) -> str:
    for encoding in ("utf-8", "gb18030", "big5", "latin-1"):
        try:
            return clean_text(path.read_text(encoding=encoding, errors="ignore"))
        except UnicodeError:
            continue
    return clean_text(path.read_text(errors="ignore"))


def output_path(base: Path, relative_path: str) -> Path:
    safe = relative_path.replace("\\", "/")
    target = base / safe
    return target.with_suffix(target.suffix + ".txt")


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract text from planned local resources.")
    parser.add_argument("--plan", type=Path, default=DEFAULT_PLAN)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--limit", type=int, default=0, help="Maximum files to process; 0 means all.")
    parser.add_argument("--top-dir", help="Only process files under this top-level directory.")
    parser.add_argument("--max-pdf-pages", type=int, default=80)
    parser.add_argument("--max-xlsx-rows", type=int, default=500)
    args = parser.parse_args()

    if not args.plan.exists():
        raise SystemExit("Missing distillation plan. Run scripts/build_distillation_plan.py first.")

    processed = 0
    skipped = 0
    errors = 0
    args.output_dir.mkdir(parents=True, exist_ok=True)
    report_path = args.output_dir.parent / "extract-report.jsonl"

    with args.plan.open("r", encoding="utf-8") as plan, report_path.open("w", encoding="utf-8", newline="\n") as report:
        for line in plan:
            item = json.loads(line)
            if item.get("action") != "extract-text":
                continue
            if args.top_dir and item.get("top_dir") != args.top_dir:
                continue
            if args.limit and processed >= args.limit:
                break
            source = Path(str(item["path"]))
            rel = str(item["relative_path"])
            target = output_path(args.output_dir, rel)
            try:
                ext = str(item.get("extension", "")).lower()
                if ext == ".pdf":
                    text = extract_pdf(source, args.max_pdf_pages)
                elif ext == ".docx":
                    text = extract_docx(source)
                elif ext == ".xlsx":
                    text = extract_xlsx(source, args.max_xlsx_rows)
                else:
                    text = extract_plain(source)
                if not text:
                    skipped += 1
                    status = "empty"
                else:
                    target.parent.mkdir(parents=True, exist_ok=True)
                    header = f"# {rel}\n\nSource: `{source}`\n\n"
                    target.write_text(header + text + "\n", encoding="utf-8", newline="\n")
                    processed += 1
                    status = "ok"
                report.write(json.dumps({"status": status, "relative_path": rel, "target": str(target)}, ensure_ascii=False) + "\n")
            except Exception as exc:  # noqa: BLE001 - report and continue batch extraction.
                errors += 1
                report.write(json.dumps({"status": "error", "relative_path": rel, "error": repr(exc)}, ensure_ascii=False) + "\n")

    print(f"processed={processed} skipped={skipped} errors={errors}")
    print(f"report={report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
